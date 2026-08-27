#!/usr/bin/env python3
"""
验证生成的 .docx 文档 — 检查修订痕迹和批注逻辑

用法：
  python3 scripts/verify_mock_docx.py output/金竹山矿业_行政处罚决定书_模拟_20260808.docx
  python3 scripts/verify_mock_docx.py output/金竹山火力发电_处罚决定书.docx
"""
import argparse
import sys
from docx import Document
from docx.shared import RGBColor

TERRA = RGBColor(0xC9, 0x7C, 0x3E)   # 陶土橙 — AI
GRAY_BLUE = RGBColor(0x5B, 0x6C, 0x85)  # 灰蓝 — 批注
RED = RGBColor(0xCC, 0x3B, 0x3B)
GREEN = RGBColor(0x5A, 0x7D, 0x3C)


def verify(filepath):
    print(f'📄 打开文档: {filepath}')
    doc = Document(filepath)

    total = len(doc.paragraphs)
    ai_paras = []
    comment_lines = []
    sections = set()

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # 检测章标题
        if text.startswith(('一、', '二、', '三、', '四、', '五、')):
            sections.add(text[:30])

        # 检测 AI 标记段落（陶土橙色文字）
        for run in p.runs:
            if (
                hasattr(run, 'font')
                and run.font.color
                and run.font.color.rgb == TERRA
            ):
                ai_paras.append({
                    'paragraph_num': i,
                    'preview': text[:100].replace('\n', ' '),
                    'full': text,
                    'has_underline': run.underline,
                })
                break

        # 检测批注行（灰蓝色）
        for run in p.runs:
            if (
                hasattr(run, 'font')
                and run.font.color
                and run.font.color.rgb == GRAY_BLUE
                and run.font.size
                and run.font.size.pt <= 10
            ):
                comment_lines.append({
                    'paragraph_num': i,
                    'text': text[:120],
                })
                break

    # ── 报告 ────────────────────────────────
    print(f'\n{"="*60}')
    print(f'✅ 验证完成')
    print(f'{"="*60}')

    print(f'\n📊 基础统计:')
    print(f'  总段落数: {total}')
    print(f'  章标题: {len(sections)} 个')
    for s in sorted(sections):
        print(f'    {s}')

    print(f'\n🟠 AI 标记段落 (陶土橙 #C97C3E): {len(ai_paras)} 处')
    for i, a in enumerate(ai_paras, 1):
        ul = '有下划线' if a['has_underline'] else '无下划线'
        print(f'  {i}. P{a["paragraph_num"]} [{ul}] {a["preview"]}...')

    print(f'\n💬 批注/引用 (灰蓝 #5B6C85): {len(comment_lines)} 条')
    for i, c in enumerate(comment_lines, 1):
        print(f'  {i}. P{c["paragraph_num"]} {c["text"]}')

    # ── 判定 ────────────────────────────────
    issues = []
    if not ai_paras:
        issues.append('❌ 未检测到任何 AI 标记段落（陶土橙色文字）')
    if not comment_lines:
        issues.append('⚠️  未检测到批注记录（灰蓝小字）')

    if issues:
        print(f'\n{"!"*60}')
        for iss in issues:
            print(f'  {iss}')
        print(f'{"!"*60}')
    else:
        print(f'\n🎉 所有验证项通过！')

    return {
        'paragraphs': total,
        'ai_marked': len(ai_paras),
        'comments': len(comment_lines),
        'sections': len(sections),
        'issues': issues,
    }


def main():
    parser = argparse.ArgumentParser(description='验证生成的 .docx 文档')
    parser.add_argument('file', help='.docx 文件路径')
    args = parser.parse_args()

    try:
        result = verify(args.file)
        if result['issues']:
            sys.exit(1)
    except Exception as e:
        print(f'❌ 验证异常: {e}', file=sys.stderr)
        sys.exit(2)


if __name__ == '__main__':
    main()
