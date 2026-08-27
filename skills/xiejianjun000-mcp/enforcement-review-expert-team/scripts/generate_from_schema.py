#!/usr/bin/env python3
"""
全流程文书生成：Schema YAML + 案件 JSON → .docx
==============================================
读取 38_行政处罚决定书.yaml 的字段定义和分支逻辑，
自动从案件 JSON 填充，处理勾选、分支决策、格式整理，
输出带修订痕迹（陶土橙 AI 标记）的 .docx 文件。

用法：
  python3 scripts/generate_from_schema.py \\
    --schema eco-enforcement-review-team/skills/doc-panel/schemas/38_行政处罚决定书.yaml \\
    --case data/cases/金竹山火力发电.json \\
    --output output/金竹山火力发电_处罚决定书.docx
"""
import argparse
import json
import logging
import os
import sys
import yaml
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 日志 ──────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S',
)
log = logging.getLogger('doc-generator')

# ── 颜色常量 ───────────────────────────────────
TERRA = RGBColor(0xC9, 0x7C, 0x3E)      # 陶土橙 — AI 生成标记
GRAY_BLUE = RGBColor(0x5B, 0x6C, 0x85)  # 灰蓝 — 引用/批注
RED = RGBColor(0xCC, 0x3B, 0x3B)
GREEN = RGBColor(0x5A, 0x7D, 0x3C)
BLACK = RGBColor(0x2D, 0x2D, 0x2D)

# ── 工具函数 ───────────────────────────────────

def deep_get(dct, path, default=None):
    """从嵌套字典取 'a.b.c' 路径值"""
    keys = path.split('.')
    current = dct
    for k in keys:
        if isinstance(current, dict):
            current = current.get(k)
        else:
            return default
        if current is None:
            return default
    return current


def add_run(para, text, **kwargs):
    """添加格式文本运行"""
    run = para.add_run(text)
    font = run.font
    font.size = kwargs.get('size', Pt(14))
    font.color.rgb = kwargs.get('color', BLACK)
    font.bold = kwargs.get('bold', False)
    if kwargs.get('underline'):
        run.underline = True
    if kwargs.get('italic'):
        font.italic = True
    return run


def add_ai_insertion(para, text, source=''):
    """AI 生成内容：陶土橙下划线 + 浅沙底纹"""
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = TERRA
    run.underline = True
    # 浅沙底纹
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFF5EE')
    shd.set(qn('w:val'), 'clear')
    run._r.get_or_add_rPr().append(shd)
    # 来源标记
    if source:
        src = para.add_run(f' [{source}]')
        src.font.size = Pt(8)
        src.font.color.rgb = GRAY_BLUE
        src.font.italic = True
    return run


def add_comment_ref(para, text):
    """灰蓝批注引用"""
    run = para.add_run(f' [{text}]')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_BLUE
    return run


def add_ai_label(para):
    """'AI 生成' 标签"""
    run = para.add_run('  🤖 AI 生成')
    run.font.size = Pt(8)
    run.font.color.rgb = TERRA
    run.font.italic = True
    return run


def fill_date_field(date_val):
    """统一日期格式为'Y年m月d日'"""
    if not date_val:
        return '____年__月__日'
    if isinstance(date_val, date):
        return date_val.strftime('%Y年%-m月%-d日')
    if isinstance(date_val, str):
        try:
            d = date.fromisoformat(date_val)
            return d.strftime('%Y年%-m月%-d日')
        except (ValueError, TypeError):
            return date_val
    return str(date_val)


# ── 主流程 ─────────────────────────────────────

def load_config(case_path, schema_path):
    """加载案件数据和 Schema"""
    log.info(f'加载案件数据: {case_path}')
    with open(case_path, 'r', encoding='utf-8') as f:
        case = json.load(f)
    log.info(f'  -> 当事人: {deep_get(case, "案件数据.当事人.名称")}')
    log.info(f'  -> 调查日期: {fill_date_field(deep_get(case, "案件数据.调查日期"))}')

    log.info(f'加载 Schema: {schema_path}')
    with open(schema_path, 'r', encoding='utf-8') as f:
        schema = yaml.safe_load(f)
    log.info(f'  -> 文书名称: {schema["name"]}')
    log.info(f'  -> 字段数: {len(schema["fields"])}')
    log.info(f'  -> 分支数: {len(schema.get("branches", []))}')

    return case, schema


def resolve_field(field_def, case):
    """解析单个字段值"""
    field_id = field_def['id']
    source = field_def['source']
    required = field_def.get('required', False)
    field_type = field_def.get('type', 'string')

    log.debug(f'  解析字段: [{field_id}] source={source} required={required}')

    if source.startswith('系统配置.'):
        value = deep_get(case, source)
        log.debug(f'    => 系统配置: {value}')
        return value
    elif source.startswith('系统数据.'):
        value = deep_get(case, source)
        log.debug(f'    => 系统数据: {value}')
        return value
    elif source.startswith('系统生成.'):
        value = deep_get(case, source)
        log.debug(f'    => 系统生成: {value}')
        return value
    elif source.startswith('案件数据.'):
        data_path = source.replace('案件数据.', '')
        case_data = case.get('案件数据', {})

        if '证据' in data_path:
            return {'type': 'evidence', 'data': case_data.get('证据', [])}
        elif '法律依据.违法条款' in data_path:
            return case_data.get('法律依据', {}).get('违法条款')
        elif '法律依据.处罚条款' in data_path:
            return case_data.get('法律依据', {}).get('处罚条款')
        elif '陈述申辩.' in data_path:
            return case_data.get('陈述申辩', {})
        elif '听证.' in data_path:
            return case_data.get('听证', {})
        elif '处罚决定.' in data_path:
            return case_data.get('处罚决定', {})
        elif '集体讨论.' in data_path:
            return case_data.get('集体讨论', {})
        elif '告知书.' in data_path:
            return case_data.get('告知书', {})
        else:
            value = deep_get(case_data, data_path)
            log.debug(f'    => 案件数据: {value}')
            return value
    elif source.startswith('AI生成.'):
        return {'type': 'ai_generated', 'source': source, 'def': field_def}
    else:
        log.warning(f'  未知来源: {source}')
        return None


def resolve_branch(schema, case):
    """解析分支决策"""
    branches = schema.get('branches', [])
    case_data = case.get('案件数据', {})

    has_hearing_right = case_data.get('听证权判定', False)
    has_statement = case_data.get('陈述申辩', {}).get('状态') == '已进行'
    has_hearing_requested = case_data.get('听证', {}).get('申请状态') == '已申请'
    has_fine = case_data.get('处罚决定', {}).get('是否罚款', False)

    log.info(f'分支决策参数: 听证权={has_hearing_right} 陈述申辩={has_statement} 申请听证={has_hearing_requested} 罚款={has_fine}')

    results = {}
    for branch in branches:
        bid = branch['id']
        if bid == '陈述申辩听证段':
            if has_hearing_right and not has_statement and not has_hearing_requested:
                selected = 'branch_no_action_hearing_right'
            elif has_hearing_right and not has_statement and has_hearing_requested:
                selected = 'branch_hearing_only'
            elif has_hearing_right and has_statement and not has_hearing_requested:
                selected = 'branch_statement_only_hearing_right'
            elif has_hearing_right and has_statement and has_hearing_requested:
                selected = 'branch_both_hearing_right'
            elif not has_hearing_right and not has_statement:
                selected = 'branch_no_statement_no_hearing'
            elif not has_hearing_right and has_statement:
                selected = 'branch_statement_only_no_hearing'
            results[bid] = selected
            log.info(f'  陈述申辩听证段 -> {selected}')
        elif bid == '罚款收缴方式段':
            if has_fine:
                selected = 'branch_include_payment'
            else:
                selected = 'branch_no_payment'
            results[bid] = selected
            log.info(f'  罚款收缴方式段 -> {selected}')
    return results


def build_evidence_text(evidence_list):
    """生成证据列表文本"""
    if not evidence_list:
        return '证据材料见案卷。'
    lines = []
    for i, ev in enumerate(evidence_list, 1):
        line = (
            f'{i}. 证据名称：{ev.get("名称", "")}；'
            f'提取（作出）时间：{ev.get("提取时间", "")}；'
            f'提供（作出）单位：{ev.get("提供单位", "")}；'
            f'证明内容：{ev.get("证明内容", "")}。'
        )
        lines.append(line)
    return lines


def generate_ai_text(def_dict, case):
    """AI 生成字段的模拟逻辑"""
    field_id = def_dict['id']
    case_data = case.get('案件数据', {})

    if field_id == '违法事实陈述':
        facts = case_data.get('违法事实', {})
        return (
            f'我厅（局）于{facts.get("时间", "__年__月__日")}对你（单位）进行了调查，'
            f'发现你（单位）实施了以下生态环境违法行为：\n'
            f'{facts.get("行为", "")}。\n'
            f'经查，{facts.get("情节", "")}。'
        )

    if field_id == '证据列表':
        # 证据列表由 build_evidence_text 处理，这里返回特殊标记
        return '{{EVIDENCE_LIST}}'

    if field_id == '违反的法律规定':
        law = case_data.get('法律依据', {}).get('违法条款', '')
        return f'你（单位）的上述行为违反了{law}的规定。'

    if field_id == '陈述申辩采纳情况':
        review = case_data.get('复核意见', {})
        if review.get('采纳情况') == '部分采纳':
            return f'对当事人提出的事实、理由和证据{review["采纳情况"]}。{review["理由"]}'
        elif review.get('采纳情况') == '全部采纳':
            return f'对当事人提出的事实、理由和证据全部采纳。{review.get("理由", "")}'
        return f'对当事人提出的事实、理由和证据不予采纳。{review.get("理由", "")}'

    if field_id == '处罚依据':
        law = case_data.get('法律依据', {}).get('处罚条款', '')
        return f'依据{law}之规定'

    if field_id == '裁量基准依据':
        return (
            '参照《湖南省生态环境保护行政处罚裁量权基准规定（2021版）》，'
            '综合考虑以下裁量因素：①违法行为造成的环境污染情况（CEMS数据超标涉及SO₂/NOₓ/颗粒物）；'
            '②违法行为持续时间为14天；③当事人在调查中配合态度较好。'
            '综合裁量：基准罚款50万元，从轻5%。'
        )

    if field_id == '行政处罚种类':
        penalty = case_data.get('处罚决定', {})
        parts = []
        if penalty.get('是否罚款'):
            parts.append(f'罚款（大写）{penalty.get("罚款金额", "")}')
        if penalty.get('没收违法所得金额'):
            parts.append(f'没收违法所得（大写）{penalty.get("没收违法所得金额")}')
        if penalty.get('其他处罚'):
            parts.append(penalty.get('其他处罚'))
        return '；\n'.join(parts) if parts else '警告'

    if field_id == '罚款金额':
        return case_data.get('处罚决定', {}).get('罚款金额', '')

    return f'[AI生成: {field_id}]'


def generate_document(case, schema, output_path):
    """主生成方法"""
    log.info('=== 开始生成文书 ===')

    # 解析分支
    branch_results = resolve_branch(schema, case)

    doc = Document()

    # ── 页面设置 ────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 标题 ────────────────────────────────
    dept_name = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '部门名称'), case
    )
    log.info(f'部门名称: {dept_name}')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, dept_name or 'XXX生态环境厅（局）', bold=True, size=Pt(22))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, schema['name'], bold=True, size=Pt(24))

    # 文号
    year = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '文号年份'), case
    ) or '2026'
    seq = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '文号序号'), case
    ) or '3'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f'{dept_name[:2] if dept_name else "娄"}环罚〔{year}〕{seq}号', size=Pt(16))

    doc.add_paragraph('')

    # ── 当事人信息 ──────────────────────────
    entity_name = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '当事人名称或姓名'), case
    )
    legal_rep = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '法定代表人或负责人或经营者'), case
    )
    credit_code = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '统一社会信用代码或身份证件号码'), case
    )
    address = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '地址或住址'), case
    )

    log.info(f'当事人: {entity_name}')
    log.info(f'法定代表人: {legal_rep}')
    log.info(f'证件号码: {credit_code}')
    log.info(f'地址: {address}')

    info_fields = [
        ('当事人名称', entity_name),
        ('法定代表人', legal_rep),
        ('统一社会信用代码', credit_code),
        ('地址', address),
    ]
    for label, value in info_fields:
        p = doc.add_paragraph()
        add_run(p, f'{label}：', bold=True)
        add_run(p, str(value) if value else '___________')

    doc.add_paragraph('')

    # ── 一、违法事实和证据 ────────────────────
    p = doc.add_paragraph()
    add_run(p, '一、生态环境违法事实和证据', bold=True)
    doc.add_paragraph('')

    # 关联法定性：提示处罚金额超过50万 + 停产整治 → 需听证权
    facts_ai = generate_ai_text(
        next(f for f in schema['fields'] if f['id'] == '违法事实陈述'), case
    )
    log.info(f'违法事实长度: {len(facts_ai)}字符')
    p = doc.add_paragraph()
    add_ai_insertion(p, facts_ai.replace('/n', '\n'), source='数据芯')
    add_ai_label(p)
    doc.add_paragraph('')

    # 证据
    evidence = case.get('案件数据', {}).get('证据', [])
    p = doc.add_paragraph()
    add_run(p, '以上事实，有以下主要证据证明：', bold=True)
    doc.add_paragraph('')

    evidence_lines = build_evidence_text(evidence)
    log.info(f'证据条数: {len(evidence_lines)}')
    for line in evidence_lines:
        p = doc.add_paragraph()
        add_ai_insertion(p, line)
    add_ai_label(doc.add_paragraph())

    doc.add_paragraph('')

    # 违法条款
    law_text = generate_ai_text(
        next(f for f in schema['fields'] if f['id'] == '违反的法律规定'), case
    )
    p = doc.add_paragraph()
    add_ai_insertion(p, law_text, source='法条通')
    add_ai_label(p)

    doc.add_paragraph('')

    # ── 二、陈述申辩 ────────────────────────
    p = doc.add_paragraph()
    add_run(p, '二、陈述、申辩等权利内容的采纳情况及理由', bold=True)
    doc.add_paragraph('')

    # 告知书
    tell_year = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '告知书文号年份'), case
    ) or '2026'
    tell_seq = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '告知书文号序号'), case
    ) or '3'
    tell_date_val = resolve_field(
        next(f for f in schema['fields'] if f['id'] == '告知书送达日期'), case
    )
    tell_date = fill_date_field(tell_date_val)

    p = doc.add_paragraph()
    add_run(p, f'我厅（局）于 {tell_date} 以《行政处罚事先（听证）告知书》')
    add_run(p, f'（{dept_name[:2] if dept_name else ""}环罚告〔{tell_year}〕{tell_seq}号）')
    add_run(p, '告知你（单位）陈述申辩权、听证权。')
    doc.add_paragraph('')

    branch_selected = branch_results.get('陈述申辩听证段', '')
    log.info(f'陈述申辩分支: {branch_selected}')

    # 根据分支渲染
    has_hearing = case.get('案件数据', {}).get('听证权判定', False)
    has_stmnt = case.get('案件数据', {}).get('陈述申辩', {}).get('状态') == '已进行'

    if branch_selected == 'branch_statement_only_hearing_right':
        p = doc.add_paragraph()
        add_run(p, '☑ ', bold=True, color=RED)
        add_run(p, '你（单位）在法定期限内未要求听证，视为放弃要求听证的权利。')

        stmnt_date = fill_date_field(
            case.get('案件数据', {}).get('陈述申辩', {}).get('日期')
        )
        stmnt_content = case.get('案件数据', {}).get('陈述申辩', {}).get('内容', '')
        decision_text = generate_ai_text(
            next(f for f in schema['fields'] if f['id'] == '陈述申辩采纳情况'), case
        )

        doc.add_paragraph('')
        p = doc.add_paragraph()
        add_run(p, f'你（单位）于 {stmnt_date} 进行了陈述和申辩，提出以下意见：')

        p = doc.add_paragraph()
        add_ai_insertion(p, f'（{stmnt_content}）', source='文书成')
        add_ai_label(p)

        doc.add_paragraph('')
        p = doc.add_paragraph()
        add_run(p, '我厅（局）充分听取后复核认为，')
        add_ai_insertion(p, decision_text, source='文书成')
        add_ai_label(p)

    elif branch_selected == 'branch_no_action_hearing_right':
        p = doc.add_paragraph()
        add_run(p, '☑ ', bold=True, color=RED)
        add_run(p, '你（单位）在期限内未进行陈述和申辩，也未要求听证，视为放弃陈述和申辩、要求听证的权利。')

    doc.add_paragraph('')

    # ── 三、处罚依据 ────────────────────────
    p = doc.add_paragraph()
    add_run(p, '三、行政处罚的依据、种类，以及裁量基准运用的理由和依据', bold=True)
    doc.add_paragraph('')

    penalty_basis = generate_ai_text(
        next(f for f in schema['fields'] if f['id'] == '处罚依据'), case
    )
    discretion = generate_ai_text(
        next(f for f in schema['fields'] if f['id'] == '裁量基准依据'), case
    )
    penalty_type = generate_ai_text(
        next(f for f in schema['fields'] if f['id'] == '行政处罚种类'), case
    )

    log.info(f'处罚依据: {penalty_basis[:60]}...')

    p = doc.add_paragraph()
    add_ai_insertion(p, f'{penalty_basis}，{discretion}', source='法条通')
    add_ai_label(p)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '我厅（局）决定对你（单位）处以如下行政处罚：', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_ai_insertion(p, penalty_type)
    add_ai_label(p)

    doc.add_paragraph('')

    # ── 四、履行方式 ────────────────────────
    p = doc.add_paragraph()
    add_run(p, '四、行政处罚的履行方式和期限', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '限于接到本处罚决定之日起十五日内到指定的银行或者通过电子支付系统缴纳罚款。逾期不缴纳罚款的，我厅（局）可以根据《中华人民共和国行政处罚法》第七十二条第一款第一项之规定每日按罚款数额的百分之三加处罚款。')
    doc.add_paragraph('')

    # 收款信息
    bank = resolve_field(next(f for f in schema['fields'] if f['id'] == '收款银行'), case)
    acct_name = resolve_field(next(f for f in schema['fields'] if f['id'] == '收款户名'), case)
    acct_no = resolve_field(next(f for f in schema['fields'] if f['id'] == '收款账号'), case)

    for label, val in [('收款银行', bank), ('户名', acct_name), ('账号', acct_no)]:
        p = doc.add_paragraph()
        add_run(p, f'{label}：', bold=True)
        add_run(p, str(val) if val else '___________')

    doc.add_paragraph('')

    # ── 五、救济途径 ────────────────────────
    p = doc.add_paragraph()
    add_run(p, '五、申请行政复议或者提起行政诉讼的途径和期限', bold=True)
    doc.add_paragraph('')

    复议机关 = resolve_field(next(f for f in schema['fields'] if f['id'] == '行政复议机关'), case)
    法院 = resolve_field(next(f for f in schema['fields'] if f['id'] == '行政诉讼法院'), case)

    p = doc.add_paragraph()
    add_run(p, f'你（单位）如不服本处罚决定，可在收到本处罚决定书之日起六十日内向{复议机关}申请行政复议，也可以在六个月内向{法院}提起行政诉讼。申请行政复议或者提起行政诉讼，不停止行政处罚决定的执行。')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '逾期不申请行政复议，不提起行政诉讼，又不履行本处罚决定，我厅（局）将依法申请人民法院强制执行。')

    # ── 落款 ────────────────────────────────
    doc.add_paragraph('')
    doc.add_paragraph('')
    decision_date = resolve_field(next(f for f in schema['fields'] if f['id'] == '作出日期'), case)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f'{dept_name}（印章）', size=Pt(16))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, fill_date_field(decision_date), size=Pt(16))

    # ── 批注记录 ────────────────────────────
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    add_run(p, '═' * 40, color=GRAY_BLUE, size=Pt(10))
    p = doc.add_paragraph()
    add_run(p, '【审阅批注】（模拟右侧栏 Office 文书协同）', color=GRAY_BLUE, size=Pt(10), bold=True)

    annotations = [
        ('文书成（AI）', '08:15', "违法事实陈述根据CEMS数据台账生成。24次超标=24次'启动'标记=1:1对应。"),
        ('李建国', '08:22', '确认数据准确。改为"涉嫌伪造自动监测数据"的措辞。'),
        ('文书成（AI）', '08:23', '已修改。调整为"涉嫌通过工况标记造假规避超标记录"。✓'),
        ('李建国', '08:25', f'罚款{case.get("案件数据", {}).get("处罚决定", {}).get("罚款金额", "")}需在决定书中单独列出。'),
        ('文书成（AI）', '08:26', '已在第三部分"罚款（大写）"行独立显示。'),
        ('李建国', '08:30', '整体审核通过。生成终稿。'),
    ]
    for author, time, text in annotations:
        p = doc.add_paragraph()
        rc = TERRA if 'AI' in author else GRAY_BLUE
        add_run(p, f'[{time}] ', color=GRAY_BLUE, size=Pt(9))
        add_run(p, f'{author}：', color=rc, size=Pt(9), bold=True)
        add_run(p, text, color=GRAY_BLUE, size=Pt(9))

    # ── 保存 ────────────────────────────────
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    # ── 统计 ────────────────────────────────
    paragraphs_count = len(doc.paragraphs)
    ai_count = sum(1 for p in doc.paragraphs if any(
        hasattr(r, 'font') and r.font.color.rgb == TERRA for r in p.runs
    ))
    log.info(f'=== 生成完成 ===')
    log.info(f'  输出文件: {output_path}')
    log.info(f'  总段落数: {paragraphs_count}')
    log.info(f'  AI 标记段: {ai_count}')
    log.info(f'  分支决策: {branch_results}')
    log.info(f'  批注记录: {len(annotations)} 条')

    return output_path


def main():
    parser = argparse.ArgumentParser(description='全流程文书生成')
    parser.add_argument('--schema', required=True, help='Schema YAML 路径')
    parser.add_argument('--case', required=True, help='案件 JSON 路径')
    parser.add_argument('--output', default='output/output.docx', help='输出 .docx 路径')
    args = parser.parse_args()

    try:
        case, schema = load_config(args.case, args.schema)
        output = generate_document(case, schema, args.output)
        print(f'✅ 文书已生成: {output}')
    except Exception as e:
        log.exception(f'生成失败: {e}')
        sys.exit(1)


if __name__ == '__main__':
    main()
