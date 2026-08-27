#!/usr/bin/env python3
"""
生成模拟 .docx 文书 — 行政处罚决定书（金竹山矿业 CEMS 超标案）
验证：修订痕迹 + 陶土橙 AI 标记 + 批注 + 人机同步效果
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime, logging, time

# ── 日志配置 ──────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S',
)
log = logging.getLogger('mock-doc-generator')

# ── 颜色常量 ──────────────────────────────────
TERRA = RGBColor(0xC9, 0x7C, 0x3E)      # 陶土橙 — AI 修改标记
GRAY_BLUE = RGBColor(0x5B, 0x6C, 0x85)  # 灰蓝 — 引用/批注
RED = RGBColor(0xCC, 0x3B, 0x3B)
OLIVE = RGBColor(0x5A, 0x7D, 0x3C)
BLACK = RGBColor(0x2D, 0x2D, 0x2D)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'output')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── 工具函数 ──────────────────────────────────

def add_run(para, text, bold=False, color=BLACK, size=Pt(12), underline=False, ul_color=None):
    """添加带格式的文本运行"""
    run = para.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    if underline:
        run.underline = True
    return run

def add_ai_insertion(para, text, size=Pt(12)):
    """添加 AI 插入的内容 — 陶土橙下划线 + 灰蓝底"""
    run = para.add_run(text)
    run.font.size = size
    run.font.color.rgb = TERRA
    run.underline = True
    # 添加底纹（模拟修订痕迹中 AI 内容高亮）
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FDF5ED')
    shading.set(qn('w:val'), 'clear')
    run._r.get_or_add_rPr().append(shading)
    return run

def add_comment_ref(para, text):
    """添加批注引用 — 灰蓝色小字"""
    run = para.add_run(f" [{text}]")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_BLUE
    return run

def add_ai_label(para):
    """标注"AI 生成"标签"""
    run = para.add_run(" 🤖 AI 生成 · 可撤销")
    run.font.size = Pt(8)
    run.font.color.rgb = TERRA
    run.font.italic = True
    return run

# ── 主文档生成 ────────────────────────────────

def generate():
    t0 = time.time()
    log.info('=== 开始生成模拟文书 ===')

    doc = Document()

    # ── 页面设置 ──────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(16)
    log.info('页面设置: A4 21×29.7cm, 仿宋 16pt')

    # ══════════════════════════════════════════
    # 标题
    # ══════════════════════════════════════════
    log.info('生成标题...')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '娄底市生态环境局', bold=True, size=Pt(22))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '行政处罚决定书', bold=True, size=Pt(24))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '娄环罚〔2026〕3号', size=Pt(16))

    doc.add_paragraph('')  # 空行

    # ══════════════════════════════════════════
    # 当事人信息（人工填写，无 AI 标记）
    # ══════════════════════════════════════════
    log.info('写入当事人信息...')
    p = doc.add_paragraph()
    add_run(p, '当事人名称：', bold=True)
    add_run(p, '金竹山矿业有限责任公司')

    p = doc.add_paragraph()
    add_run(p, '法定代表人：', bold=True)
    add_run(p, '张建国')

    p = doc.add_paragraph()
    add_run(p, '统一社会信用代码：', bold=True)
    add_run(p, '91431381MA4L7XXXXX')

    p = doc.add_paragraph()
    add_run(p, '地址：', bold=True)
    add_run(p, '湖南省冷水江市金竹山镇')

    doc.add_paragraph('')

    # ══════════════════════════════════════════
    # 一、违法事实和证据
    # ══════════════════════════════════════════
    log.info('写入一、违法事实和证据...')
    p = doc.add_paragraph()
    add_run(p, '一、生态环境违法事实和证据', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '我厅（局）于 ')
    add_run(p, '2026年7月15日', bold=True)
    add_run(p, ' 对你（单位）进行了调查，发现你（单位）实施了以下生态环境违法行为：')
    doc.add_paragraph('')

    # —— AI 生成段落（陶土橙标记）——
    log.info('  -> AI生成: 违法事实陈述')
    p = doc.add_paragraph()
    add_ai_insertion(p, '你（单位）金竹山矿业有限责任公司所属金竹山火力发电厂，在2026年7月1日至7月14日期间，烟气连续排放监测系统（CEMS）数据显示二氧化硫排放浓度多次超过《火电厂大气污染物排放标准》（GB 13223-2011）表1规定的200mg/m³限值。')
    add_ai_label(p)
    doc.add_paragraph('')

    # —— 补充（人工编辑）——
    log.info('  -> AI生成: 旁路排放嫌疑（含数据芯引用）')
    p = doc.add_paragraph()
    add_run(p, '经执法人员现场核查，')
    add_ai_insertion(p, '超标时段主要集中在夜间22:00至次日凌晨4:00，且该时段企业用电量骤降42%，与正常生产负荷明显不符，存在旁路排放嫌疑。')
    add_comment_ref(p, '数据芯 · 大数据分析')
    add_ai_label(p)
    doc.add_paragraph('')

    # —— 证据列举（人工填写格式，AI 填充内容）——
    p = doc.add_paragraph()
    add_run(p, '以上事实，有以下主要证据证明：', bold=True)
    doc.add_paragraph('')

    evidence = [
        ('营业执照、法定代表人身份证复印件各1份', '2026-07-15', '你（单位）', '证明当事人的主体资格'),
        ('《现场检查（勘察）笔录》1份', '2026-07-15', '娄底市生态环境局', '执法人员对金竹山火力发电厂进行现场检查，发现CEMS数据显示超标'),
        ('CEMS历史数据导出台账（2026.7.1-7.14）', '2026-07-15', '你（单位）', '证明14天内有19次二氧化硫小时均值超标'),
        ('企业用电量曲线（2026.7.1-7.14）', '2026-07-16', '国网冷水江供电公司', '证明夜间用电量骤降与超标时段吻合'),
        ('调查询问笔录（张建国）1份', '2026-07-16', '娄底市生态环境局', '法定代表人承认脱硫设施夜间存在运行不稳定情况'),
    ]
    log.info(f'  -> AI生成: 证据列表（{len(evidence)}条）')
    for i, (name, ev_time, source, content) in enumerate(evidence, 1):
        p = doc.add_paragraph()
        add_ai_insertion(p, f'{i}. 证据名称：{name}；提取（作出）时间：{ev_time}；提供（作出）单位：{source}；证明内容：{content}。')
    add_ai_label(doc.add_paragraph())

    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_ai_insertion(p, '你（单位）的上述行为违反了《中华人民共和国大气污染防治法》第十八条"企业事业单位和其他生产经营者建设对大气环境有影响的项目，应当依法进行环境影响评价、公开环境影响评价文件；向大气排放污染物的，应当符合大气污染物排放标准，遵守重点大气污染物排放总量控制要求"的规定。')
    add_ai_label(p)

    doc.add_paragraph('')

    # ══════════════════════════════════════════
    # 二、陈述申辩
    # ══════════════════════════════════════════
    log.info('写入二、陈述申辩听证段...')
    log.info('  分支决策: branch_statement_only_hearing_right (有听证权=是, 陈述申辩=已进行, 申请听证=未申请)')
    p = doc.add_paragraph()
    add_run(p, '二、陈述、申辩等权利内容的采纳情况及理由', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '我厅（局）于 ')
    add_run(p, '2026年7月20日', bold=True)
    add_run(p, ' 以《行政处罚事先（听证）告知书》（娄环罚告〔2026〕3号）告知你（单位）陈述申辩权、听证权。')
    doc.add_paragraph('')

    # —— 分支决策（人工选中）——
    p = doc.add_paragraph()
    add_run(p, '☑ ', bold=True, color=RED)
    add_run(p, '你（单位）于 ', bold=True)
    add_run(p, '2026年7月22日', bold=True)
    add_run(p, ' 进行了陈述和申辩，提出以下意见：')

    p = doc.add_paragraph()
    add_ai_insertion(p, '（当事人意见：承认CEMS数据超标属实，但辩称系脱硫设施老化导致，非主观故意，请求减轻处罚。并提交了《脱硫设施维修记录》和《2026年下半年脱硫设施升级改造计划》。）')
    add_ai_label(p)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '我厅（局）充分听取后复核认为，')
    add_ai_insertion(p, '对当事人提出的事实、理由和证据部分采纳。脱硫设施老化属于企业自身设备维护责任，不构成减轻处罚的法定事由；但当事人积极配合调查、主动提交整改计划，在裁量时可酌情考虑。')
    add_ai_label(p)

    doc.add_paragraph('')

    # ══════════════════════════════════════════
    # 三、处罚依据和种类
    # ══════════════════════════════════════════
    log.info('写入三、处罚依据和裁量...')
    p = doc.add_paragraph()
    add_run(p, '三、行政处罚的依据、种类，以及裁量基准运用的理由和依据', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '依据《中华人民共和国大气污染防治法》第九十九条第二项"超过大气污染物排放标准或者超过重点大气污染物排放总量控制指标排放大气污染物的，由县级以上人民政府生态环境主管部门责令改正或者限制生产、停产整治，并处十万元以上一百万元以下的罚款"之规定，参照《湖南省生态环境保护行政处罚裁量权基准规定（2021版）》中"排放大气污染物超过排放标准"的裁量基准：')
    doc.add_paragraph('')

    # —— 裁量说明（AI 生成）——
    p = doc.add_paragraph()
    add_ai_insertion(p, '裁量因素分析：①排放超标持续时间14天，属于"持续时间较长（7天以上）"档；②超标时段集中在夜间，存在规避监管嫌疑，从重10%；③企业积极配合调查并提交整改计划，从轻5%。综合裁量：基准罚款30万元 + 从重10%（3万元）- 从轻5%（1.5万元）= 31.5万元。大写：叁拾壹万伍仟元整。')
    add_ai_label(p)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '我厅（局）决定对你（单位）处以如下行政处罚：', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '☑ 罚款（大写）', bold=True, color=RED)
    add_ai_insertion(p, '叁拾壹万伍仟元整')
    add_ai_label(p)

    doc.add_paragraph('')

    # ══════════════════════════════════════════
    # 四、履行方式
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    add_run(p, '四、行政处罚的履行方式和期限', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '限于接到本处罚决定之日起十五日内到指定的银行或者通过电子支付系统缴纳罚款。逾期不缴纳罚款的，我厅（局）可以根据《中华人民共和国行政处罚法》第七十二条第一款第一项之规定每日按罚款数额的百分之三加处罚款。')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '收款银行：', bold=True)
    add_run(p, '中国工商银行冷水江支行')
    p = doc.add_paragraph()
    add_run(p, '户名：', bold=True)
    add_run(p, '娄底市财政局非税收入汇缴结算户')
    p = doc.add_paragraph()
    add_run(p, '账号：', bold=True)
    add_run(p, '1901024209200XXXXXX')

    doc.add_paragraph('')

    # ══════════════════════════════════════════
    # 五、救济途径
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    add_run(p, '五、申请行政复议或者提起行政诉讼的途径和期限', bold=True)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '你（单位）如不服本处罚决定，可在收到本处罚决定书之日起六十日内向娄底市人民政府申请行政复议，也可以在六个月内向娄底市娄星区人民法院提起行政诉讼。申请行政复议或者提起行政诉讼，不停止行政处罚决定的执行。')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '逾期不申请行政复议，不提起行政诉讼，又不履行本处罚决定，我厅（局）将依法申请人民法院强制执行。')

    # ── 落款 ──────────────────────────────────
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, '娄底市生态环境局（印章）', size=Pt(16))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f'{datetime.date.today().strftime("%Y年%m月%d日")}', size=Pt(16))

    # ── 底部批注区（模拟右侧栏审阅模式）──
    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    add_run(p, '═' * 40, color=GRAY_BLUE, size=Pt(10))
    p = doc.add_paragraph()
    add_run(p, '【审阅批注记录】（演示右侧栏协作）', color=GRAY_BLUE, size=Pt(10), bold=True)

    comments = [
        ('文书成（AI）', '2026-07-25 10:24', '违法事实陈述已根据CEMS数据台账自动生成，请核对超标天数。'),
        ('李建国（人工）', '2026-07-25 10:31', '超标天数确认14天无误，但第3段"旁路排放嫌疑"措辞偏重，建议改为"运行异常"。'),
        ('文书成（AI）', '2026-07-25 10:32', '已修改。当前表述调整为"与正常生产负荷不符，存在运行异常情况"。✓ 已解决'),
        ('李建国（人工）', '2026-07-25 10:35', '罚款金额复核：基准30万+从重10%-从轻5%=31.5万，计算正确。通过。'),
    ]
    for author, atime, text in comments:
        p = doc.add_paragraph()
        role_color = TERRA if 'AI' in author else GRAY_BLUE
        add_run(p, f'[{atime}] ', color=GRAY_BLUE, size=Pt(9))
        add_run(p, f'{author}：', color=role_color, size=Pt(9), bold=True)
        add_run(p, text, color=GRAY_BLUE, size=Pt(9))

    # ── 保存 ──────────────────────────────────
    filename = f'金竹山矿业_行政处罚决定书_模拟_{datetime.date.today().strftime("%Y%m%d")}.docx'
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    elapsed = time.time() - t0
    ai_paras = sum(1 for p in doc.paragraphs if any(
        hasattr(r, 'font') and r.font.color.rgb == TERRA for r in p.runs
    ))
    log.info(f'=== 生成完成 ({elapsed:.1f}s) ===')
    log.info(f'  输出文件: {filepath}')
    log.info(f'  总段落数: {len(doc.paragraphs)}')
    log.info(f'  陶土橙 AI 标记段落: {ai_paras}')
    log.info(f'  批注条数: {len(comments)}')
    print(f'✅ 文书已生成: {filepath}')
    print(f'   - 陶土橙 AI 标记: 违法事实/证据/裁量分析/处罚金额')
    print(f'   - 分支决策: 陈述申辩段手动选中')
    print(f'   - 底部批注: {len(comments)}条人机审阅记录')
    return filepath

if __name__ == '__main__':
    generate()
