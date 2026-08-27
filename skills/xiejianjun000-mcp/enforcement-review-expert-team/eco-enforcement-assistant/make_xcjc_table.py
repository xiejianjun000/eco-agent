# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
# 页边距收窄，确保一页
for s in doc.sections:
    s.top_margin = Pt(50); s.bottom_margin = Pt(50)
    s.left_margin = Pt(60); s.right_margin = Pt(60)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cell_font(cell, size=9, bold=False):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.name = '宋体'; r.font.size = Pt(size); r.bold = bold
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def shade(cell, color='D9E2F3'):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),color)
    tcPr.append(sh)

def add_row(table, labels_values, merge_from=None, bold=False, shade_color=None):
    """labels_values: list of (text) per column; merge_from: index to merge start (merge rest)."""
    cells = table.add_row().cells
    n = len(labels_values)
    for i, txt in enumerate(labels_values):
        cells[i].text = txt
        set_cell_font(cells[i], bold=bold)
        if shade_color: shade(cells[i], shade_color)
    if merge_from is not None:
        merged = cells[merge_from].merge(cells[-1])
        set_cell_font(merged, bold=bold)
    return cells

# 标题
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('娄底市生态环境局现场监察记录')
tr.bold = True; tr.font.size = Pt(14); tr.font.name='宋体'
tr.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
tp.paragraph_format.space_after = Pt(4)

# 表格
t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'

# 主体信息
add_row(t, ['被检查单位名称','冷水江锑都环保有限责任公司','排污许可证号',''])
add_row(t, ['统一社会信用代码','91431381595469974U','法人代表姓名','欧阳剑'])
add_row(t, ['地址','锡矿山街道办事处艳山红居委会','现场负责人姓名',''])
add_row(t, ['职务','','联系电话',''])
add_row(t, ['监察内容','','',''], merge_from=1)
add_row(t, ['告知信息情况','执法人员 彭旭东 、 廖伟荣 出示执法证件，依法进行检查了解有关情况，并告知当事人申请回避等权利和协助调查等义务。当事人确认签字：','',''], merge_from=1)

# 现场监察情况 表头
add_row(t, ['现场监察情况','','',''], merge_from=0, bold=True, shade_color='D9E2F3')

add_row(t, ['生产状态','☑正常生产   □非正常生产   □其它','',''], merge_from=1)
add_row(t, ['建设项目"三同时"情况','未经环评审批的新建项目 □有 ☑无 □其它；未执行"三同时"建设项目 ☑有 □无 □其它 （注：生产工艺重大变动，未重新报批环境影响评价文件）','',''], merge_from=1)
add_row(t, ['污染设施建设、验收和运行情况','☑正常运行   □不正常运行   □其它','',''], merge_from=1)
add_row(t, ['自动监控系统情况','□未安装 □正常运行 □非正常运行 □已联网 □未联网 □已验收 □未验收','',''], merge_from=1)
add_row(t, ['在线监测数据','','',''], merge_from=1)
add_row(t, ['废水排放情况','□正常排放 ☑不正常排放（雨水收集后直排北区污水处理厂）   □其它','',''], merge_from=1)
add_row(t, ['废气排放情况','☑正常排放   □不正常排放   □其它','',''], merge_from=1)
add_row(t, ['固体废物','一般固废：□有 □暂存、转移正常   □无 □暂存、转移不正常；危险废物：□有 □暂存、转移正常   □无 □暂存、转移不正常','',''], merge_from=1)
add_row(t, ['环保管理机构、污染设施运行台账、应急预案情况','环保管理机构：□有 □无；污染设施运行台账：□有 □无；环境应急预案：□有 □无','',''], merge_from=1)

# 结论
add_row(t, ['现场监察结论','经现场检查，该单位：1、因生产工艺发生变更，锅炉已停用，因该单位属于国有资产投资，故尚未拆除；2、该单位生产工艺与原有环评发生了变更；3、该单位雨水收集后，直接排放至北区污水处理厂；4、生产过程中产生的废水不外排，循环使用。','',''], merge_from=1)

# 处理意见
add_row(t, ['处理意见及相关要求','责令该单位：1、重新修订环评报告批复；2、加强日常环境管理，确保不发生环境污染事故。','',''], merge_from=1)

# 签字
add_row(t, ['执法人员（签字）','彭旭东        廖伟荣','工作单位','娄底市生态环境局'])
add_row(t, ['被检查单位现场负责人（签字）','年    月    日','',''], merge_from=1)
add_row(t, ['记录人（签字）','年    月    日','',''], merge_from=1)

out = '/Users/mac/.qclaw/workspace-agent-6458195c/现场监察记录_冷水江锑都环保_20260718.docx'
doc.save(out)
print('SAVED', out, 'rows=', len(t.rows))
