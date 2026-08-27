# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 默认字体（中文）
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), '宋体')

def line(text='', bold=False, size=10.5, align=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), '宋体')
    return p

# 标题
line('娄底市生态环境局现场监察记录', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# 主体信息
line('被检查单位名称：冷水江锑都环保有限责任公司')
line('排污许可证号：')
line('统一社会信用代码：91431381595469974U')
line('法人代表姓名：欧阳剑')
line('地址：锡矿山街道办事处艳山红居委会')
line('现场负责人姓名：          职务：          联系电话：')
line('')
line('监察内容：')
line('')

# 告知信息情况
line('告知信息情况：')
line('执法人员 彭旭东 、 廖伟荣 出示执法证件，依法进行检查了解有关情况，并告知当事人申请回避等权利和协助调查等义务。当事人确认签字：')
line('')

# 现场监察情况
line('现场监察情况：', bold=True)
line('一、生产状态：☑正常生产   □非正常生产   □其它')
line('二、建设项目"三同时"情况：未经环评审批的新建项目 □有 ☑无 □其它；未执行"三同时"建设项目 ☑有 □无 □其它')
line('    （注：生产工艺发生重大变动，未重新报批环境影响评价文件）')
line('三、污染设施建设、验收和运行情况：☑正常运行   □不正常运行   □其它')
line('四、自动监控系统情况：□未安装 □正常运行 □非正常运行 □已联网 □未联网 □已验收 □未验收')
line('五、在线监测数据：')
line('六、废水排放情况：□正常排放 ☑不正常排放（雨水收集后直排北区污水处理厂）   □其它')
line('七、废气排放情况：☑正常排放   □不正常排放   □其它')
line('八、固体废物：')
line('    一般固废：□有 □暂存、转移正常   □无 □暂存、转移不正常')
line('    危险废物：□有 □暂存、转移正常   □无 □暂存、转移不正常')
line('九、环保管理机构、污染设施运行台账、应急预案情况：')
line('    环保管理机构：□有 □无')
line('    污染设施运行台账：□有 □无')
line('    环境应急预案：□有 □无')
line('')

# 现场监察结论（问题描述）
line('现场监察结论：', bold=True)
line('经现场检查，该单位：1、因生产工艺发生变更，锅炉已停用，因该单位属于国有资产投资，故尚未拆除；2、该单位生产工艺与原有环评发生了变更；3、该单位雨水收集后，直接排放至北区污水处理厂；4、生产过程中产生的废水不外排，循环使用。')
line('')

# 处理意见（模板原文，无法条）
line('处理意见及相关要求：', bold=True)
line('责令该单位：1、重新修订环评报告批复；2、加强日常环境管理，确保不发生环境污染事故。')
line('')

# 签字栏
line('执法人员（签字）：彭旭东        廖伟荣')
line('工作单位：娄底市生态环境局')
line('')
line('被检查单位现场负责人（签字）：')
line('                                          年    月    日')
line('')
line('记录人（签字）：')
line('                                          年    月    日')

out = '/Users/mac/.qclaw/workspace-agent-6458195c/现场监察记录_冷水江锑都环保_20260718.docx'
doc.save(out)
print('SAVED:', out)
print('paragraphs:', len(doc.paragraphs))
